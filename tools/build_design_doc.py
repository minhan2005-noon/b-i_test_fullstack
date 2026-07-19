from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/Top_VPN_Services_2026_Tai_lieu_mo_ta.docx"

INK = RGBColor(16, 32, 47)
PRIMARY = RGBColor(11, 79, 108)
HEADING = RGBColor(46, 116, 181)
HEADING_DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(95, 111, 125)
LIGHT_FILL = "F2F4F7"
ACCENT_FILL = "DFF8F4"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    set_cell_margins(table)


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_text(paragraph, text, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color or INK
    return run


def add_para(doc, text="", style=None, bold_label=None):
    p = doc.add_paragraph(style=style)
    if bold_label and text.startswith(bold_label):
        add_text(p, bold_label, bold=True)
        add_text(p, text[len(bold_label):])
    else:
        add_text(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        add_text(p, item)


def add_link_line(doc, label, url):
    p = doc.add_paragraph()
    add_text(p, label, bold=True)
    add_text(p, f" {url}", color=PRIMARY)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_keep_with_next(p)
    return p


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = PRIMARY
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = HEADING
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = HEADING
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = HEADING_DARK
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.12
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Top VPN Services 2026 - UI/UX Design Documentation")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    add_text(p, f"{label}: ", bold=True, color=PRIMARY)
    add_text(p, body)


def add_color_table(doc):
    data = [
        ("Ink", "#10202f", "Text chính, tương phản cao."),
        ("Primary", "#0b4f6c", "CTA chính, cảm giác tin cậy và bảo mật."),
        ("Accent", "#14b8a6", "Highlight phụ, trạng thái tích cực, privacy."),
        ("Warning", "#f5a524", "Điểm nhấn cho top ranking/winner."),
        ("Background", "#f6f8fb", "Nền sạch, editorial, dễ đọc."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table)
    widths = [Inches(1.4), Inches(1.25), Inches(3.7)]
    headers = ["Token", "Mã màu", "Vai trò"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_FILL)
        add_text(cell.paragraphs[0], text, bold=True)
    for row in data:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_width(cells[i], widths[i])
            add_text(cells[i].paragraphs[0], text)


def add_component_table(doc):
    data = [
        ("Button", "Radius 8px, min-height 44px, primary filled, secondary outlined."),
        ("Card", "Radius 8px, border nhẹ, shadow nhẹ, rank badge rõ."),
        ("Table", "Desktop dày thông tin; mobile horizontal scroll để giữ readability."),
        ("Pill", "Label nhỏ cho strengths và trust signals."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    widths = [Inches(1.6), Inches(4.75)]
    for i, text in enumerate(["Component", "Style"]):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_FILL)
        add_text(cell.paragraphs[0], text, bold=True)
    for row in data:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_width(cells[i], widths[i])
            add_text(cells[i].paragraphs[0], text)


def build():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(title, "Top VPN Services 2026", bold=True, color=PRIMARY)

    subtitle = doc.add_paragraph(style="Subtitle")
    add_text(
        subtitle,
        "Tài liệu mô tả research, design system, quyết định thiết kế và tư duy triển khai cho homepage so sánh VPN thị trường global.",
        color=MUTED,
    )

    add_callout(
        doc,
        "Ghi chú",
        "Toàn bộ nội dung hiển thị trên website được viết bằng tiếng Anh theo đúng yêu cầu đề bài. Tài liệu này dùng tiếng Việt để trình bày quá trình thiết kế.",
    )

    add_heading(doc, "1. UI Research", 1)
    add_heading(doc, "Nguồn 1: TechRadar - Best VPN Service 2026", 2)
    add_link_line(doc, "URL:", "https://www.techradar.com/vpn/best-vpn")
    add_para(doc, "Điểm mạnh:", bold_label="Điểm mạnh:")
    add_bullets(
        doc,
        [
            "Cấu trúc ranking rõ ràng, có câu trả lời nhanh cho lựa chọn best overall trước khi người dùng đọc sâu.",
            "Provider card kết hợp được giá, nhận xét ngắn, ưu điểm và CTA.",
            "Visual hierarchy tốt: số thứ hạng, tên provider, verdict và CTA được tách rõ.",
        ],
    )
    add_para(doc, "Điểm học được và áp dụng:", bold_label="Điểm học được và áp dụng:")
    add_bullets(
        doc,
        [
            "Đặt section Best overall trước danh sách đầy đủ để người dùng có thể ra quyết định nhanh.",
            "Mỗi provider card đều có score, price, strengths và CTA trong cùng một khối nội dung.",
        ],
    )

    add_heading(doc, "Nguồn 2: Tom's Guide - Best VPN in 2026", 2)
    add_link_line(doc, "URL:", "https://www.tomsguide.com/best-picks/best-vpn")
    add_para(doc, "Điểm mạnh:", bold_label="Điểm mạnh:")
    add_bullets(
        doc,
        [
            "Dùng các thông tin thực tế như tốc độ, server locations, giới hạn thiết bị và refund terms.",
            "Trang hỗ trợ cả hành vi scan nhanh và đọc chi tiết.",
            "CTA được đặt gần các tiêu chí mua hàng, phù hợp với nhóm user có intent cao.",
        ],
    )
    add_para(doc, "Điểm học được và áp dụng:", bold_label="Điểm học được và áp dụng:")
    add_bullets(
        doc,
        [
            "Bảng comparison dùng các cột ra quyết định: best for, score, speed, devices và price.",
            "CTA được lặp lại gần dữ liệu quan trọng thay vì chỉ đặt ở đầu trang.",
        ],
    )

    add_heading(doc, "Nguồn 3: Security.org - VPN Provider Comparison", 2)
    add_link_line(doc, "URL:", "https://www.security.org/vpn/compare/")
    add_para(doc, "Điểm mạnh:", bold_label="Điểm mạnh:")
    add_bullets(
        doc,
        [
            "Tập trung vào tiêu chí thực tế và giải thích cách người dùng nên so sánh VPN.",
            "Tạo cảm giác tin cậy vì nói về testing, nhu cầu người dùng và tiêu chí đánh giá, không chỉ deal.",
            "Layout kết hợp tốt giữa nội dung giáo dục và bảng so sánh provider.",
        ],
    )
    add_para(doc, "Điểm học được và áp dụng:", bold_label="Điểm học được và áp dụng:")
    add_bullets(
        doc,
        [
            "Thêm section methodology với các tiêu chí có trọng số.",
            "Thêm FAQ/buyer questions để giảm băn khoăn trước khi người dùng click CTA.",
        ],
    )

    add_heading(doc, "2. Design System", 1)
    add_heading(doc, "Màu sắc", 2)
    add_color_table(doc)
    add_para(
        doc,
        "Lý do chọn: Người dùng VPN quan tâm nhiều đến privacy và reliability, nên palette chính dùng nhóm xanh lam - xanh teal để gợi cảm giác bảo mật. Màu amber chỉ dùng tiết chế ở các điểm ranking để kéo sự chú ý vào lựa chọn tốt nhất mà không làm trang quá salesy.",
    )

    add_heading(doc, "Typography", 2)
    add_bullets(
        doc,
        [
            "Heading: Geist Sans, weight đậm, line-height chặt nhưng vẫn dễ đọc.",
            "Body: Geist Sans, regular/medium, kích thước 16-18px để đọc thoải mái.",
            "Hierarchy: hero headline lớn, section heading vừa phải, card heading và table label gọn hơn.",
        ],
    )
    add_para(
        doc,
        "Lý do chọn: Trang comparison cần người dùng scan nhanh. Heading lớn giúp tạo sự tự tin và nhấn mạnh chủ đề, trong khi text trong card/table được giữ gọn để thông tin có mật độ tốt.",
    )

    add_heading(doc, "Spacing", 2)
    add_bullets(
        doc,
        [
            "Sử dụng hệ spacing 8px.",
            "Gap nhỏ: 8-16px.",
            "Padding card: 18-24px.",
            "Spacing giữa section: 72-104px tùy viewport.",
        ],
    )

    add_heading(doc, "Component Style", 2)
    add_component_table(doc)

    add_heading(doc, "3. Cấu trúc Homepage", 1)
    add_numbered(
        doc,
        [
            "Sticky header: giúp người dùng truy cập nhanh các phần Rankings, Compare và Method.",
            "Hero: nêu rõ mục đích trang ngay từ đầu và đặt Compare VPNs làm CTA chính.",
            "Best overall recommendation: đặt trước full list vì người dùng tìm best VPN thường muốn có câu trả lời nhanh.",
            "Provider ranking cards: mỗi card có rank, best-use label, score, price, strengths và CTA.",
            "Comparison table: hỗ trợ người dùng ra quyết định lý tính hơn khi muốn so sánh chi tiết.",
            "Methodology: tăng trust bằng cách giải thích ranking được chấm theo tiêu chí nào.",
            "Research applied và FAQ: thể hiện tư duy thiết kế và giảm do dự trước bước click cuối.",
        ],
    )

    add_heading(doc, "4. Tư duy triển khai", 1)
    add_heading(doc, "Chia section/component", 2)
    add_para(
        doc,
        "Nếu triển khai bằng React hoặc Next.js, page có thể chia thành các component sau:",
    )
    add_bullets(
        doc,
        [
            "Header",
            "Hero",
            "TrustStrip",
            "WinnerBanner",
            "ProviderCard",
            "ProviderGrid",
            "ComparisonTable",
            "Methodology",
            "ResearchApplied",
            "FAQ",
            "Footer",
        ],
    )
    add_para(
        doc,
        "Dữ liệu provider nên được lưu trong một structured array để ranking cards và comparison table có thể dùng chung một nguồn dữ liệu, tránh nhập trùng và dễ cập nhật.",
    )

    add_heading(doc, "Responsive", 2)
    add_bullets(
        doc,
        [
            "Desktop: hero dùng 2 cột, provider cards dùng 5 cột compact, comparison table full width.",
            "Tablet: hero chuyển thành 1 cột, provider cards chuyển thành 2 cột.",
            "Mobile: cards xếp 1 cột, header wrap, comparison table dùng horizontal scroll để giữ nội dung dễ đọc.",
        ],
    )

    add_heading(doc, "Phần khó implement", 2)
    add_para(
        doc,
        "Phần khó nhất là comparison table trên mobile. Nếu ép tất cả cột vào màn hình nhỏ thì text sẽ khó đọc, nên mình chọn horizontal scroll. Trong phiên bản production, có thể cải thiện thêm bằng cách chuyển mỗi row thành mobile card, nhưng hướng đó cần thêm logic component và QA responsive kỹ hơn.",
    )

    add_heading(doc, "5. Giải thích quyết định thiết kế", 1)
    add_heading(doc, "Vì sao chọn bảng màu này?", 2)
    add_para(
        doc,
        "Nhóm màu xanh lam - teal tạo cảm giác bảo mật, công nghệ và đáng tin mà không quá lạnh. Màu amber chỉ dùng cho ranking emphasis để người dùng nhận ra top recommendation nhanh hơn.",
    )
    add_heading(doc, "Vì sao chọn typography này?", 2)
    add_para(
        doc,
        "Trang cần cảm giác editorial và trustworthy, nên mình chọn sans-serif sạch, heading weight mạnh và body text dễ đọc. Typography cũng hỗ trợ hành vi scan vì người dùng so sánh VPN bằng cách di chuyển nhanh giữa score, price và label.",
    )
    add_heading(doc, "Lấy cảm hứng từ đâu?", 2)
    add_para(
        doc,
        "Cảm hứng chính đến từ TechRadar, Tom's Guide và Security.org vì các trang này kết hợp tốt giữa editorial trust, comparison density và CTA phục vụ conversion.",
    )
    add_heading(doc, "Nếu có thêm thời gian, mình sẽ cải thiện gì?", 2)
    add_bullets(
        doc,
        [
            "Thêm affiliate disclosure và bằng chứng testing thực tế.",
            "Thêm bộ lọc theo use case như streaming, gaming, travel và privacy.",
            "Thêm logo provider nếu có quyền sử dụng asset hợp lệ.",
            "Làm thêm Figma prototype cho mobile interaction và table-to-card behavior.",
        ],
    )

    add_heading(doc, "Deliverables", 1)
    add_bullets(
        doc,
        [
            "Homepage chạy được: app/page.tsx",
            "Design system và responsive styling: app/globals.css",
            "Hero visual được generate riêng: public/vpn-dashboard-hero.png",
            "Research và giải thích thiết kế: README.md và file DOCX này",
        ],
    )

    add_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build()
